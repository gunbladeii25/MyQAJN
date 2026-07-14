"""
Agent 0 — Document Intelligence
Memproses dokumen tidak berstruktur (PDF/DOCX/Excel/TXT) dan
data API berstruktur. Menukar kandungan kualitatif kepada skor
numerik menggunakan Ollama llama3.2 (local).

Tugas:
  1. Parse dokumen → ekstrak teks dan jadual
  2. Ekstrak skor numerik secara deterministik (rule-based)
  3. Kandungan kualitatif → Ollama convert → skor 0-100
  4. Kenal pasti kod sekolah dalam dokumen
  5. [NEW] Bulk DOCX: detect MULTIPLE schools from one file (Pemeriksaan JN)
  6. [NEW] AI-powered school detection via LLM for complex documents
  7. Return structured extraction result
"""

import re
import os
import json
import hashlib
import logging
import urllib.request
from typing import Optional, List

logger = logging.getLogger(__name__)

# ── Ollama setup ──────────────────────────────────────────────────────────────
OLLAMA_URL   = os.getenv("OLLAMA_URL",   "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.2")

# ── Skala konversi kualitatif standard ───────────────────────────────────────
QUAL_SCALE = {
    "cemerlang": 92, "sangat cemerlang": 95, "sangat baik": 85, "baik": 75,
    "memuaskan": 62, "sederhana baik": 65, "sederhana": 52, "lemah": 35,
    "sangat lemah": 20, "kritikal": 10, "tidak memuaskan": 40,
    "excellent": 92, "very good": 85, "good": 75, "satisfactory": 62,
    "moderate": 52, "weak": 35, "poor": 20, "critical": 10,
    "outstanding": 95, "adequate": 60, "inadequate": 38,
}

# ── SKPMG2 Dimensi ────────────────────────────────────────────────────────────
SKPMG2_DIMENSIONS = [
    "kepimpinan", "leadership", "pengurusan organisasi", "organization",
    "program pendidikan", "kurikulum", "curriculum", "academic",
    "pengajaran", "pembelajaran", "pdpc", "teaching", "learning",
    "kemajuan murid", "student progress", "achievement",
    "keselamatan", "safety", "kemudahan", "facilities", "infrastructure",
    "disiplin", "discipline", "kokurikulum", "co-curriculum",
    "kebersihan", "hygiene", "kantin", "canteen",
]

# ── Regex patterns ─────────────────────────────────────────────────────────────
SCHOOL_CODE_RE = re.compile(
    r'\b(SK|SMK|SJK|MRSM|SBP|SKK|SEKOLAH)[A-Z\s]*\d{3,}\b', re.IGNORECASE
)
SCORE_RE = re.compile(r'(\d{1,3}(?:\.\d{1,2})?)\s*(?:/\s*100|%|markah|skor|score)', re.IGNORECASE)
HEADER_SCORE_RE = re.compile(r'([A-Za-z\s]{3,40})\s*[:\-]\s*(\d{1,3}(?:\.\d{1,2})?)')


def run(file_path: str, source_type: str, source_metadata: Optional[dict] = None) -> dict:
    """
    Proses satu dokumen. source_type: pdf | docx | xlsx | txt | api_data
    Untuk api_data, file_path adalah path ke JSON file dengan data yang dah pull.
    """
    source_metadata = source_metadata or {}

    try:
        if source_type == "api_data":
            return _process_api_data(file_path, source_metadata)
        else:
            raw_text = _parse_document(file_path, source_type)
            return _process_text(raw_text, source_type, file_path, source_metadata)
    except Exception as e:
        logger.error(f"Agent 0 error for {file_path}: {e}")
        return _error_result(str(e), file_path)


def _parse_document(file_path: str, source_type: str) -> str:
    """Extract raw text dari dokumen."""
    ext = source_type.lower()

    if ext == "pdf":
        return _parse_pdf(file_path)
    elif ext in ("docx", "doc"):
        return _parse_docx(file_path)
    elif ext in ("xlsx", "xls"):
        return _parse_excel(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()


def _parse_pdf(file_path: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except ImportError:
        return f"[PDF parsing unavailable - install PyMuPDF] File: {file_path}"
    except Exception as e:
        return f"[PDF parse error: {e}]"


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include table content
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(lines)
    except ImportError:
        return f"[DOCX parsing unavailable - install python-docx] File: {file_path}"
    except Exception as e:
        return f"[DOCX parse error: {e}]"


def _parse_excel(file_path: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"[Sheet: {sheet}]")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(v) if v is not None else "" for v in row)
                if row_text.strip():
                    lines.append(row_text)
        wb.close()
        return "\n".join(lines)
    except ImportError:
        return f"[Excel parsing unavailable - install openpyxl] File: {file_path}"
    except Exception as e:
        return f"[Excel parse error: {e}]"


def _process_text(raw_text: str, source_type: str, file_path: str, meta: dict) -> dict:
    """Process extracted text — detect scores, identify qualitative parts."""
    school_code = _extract_school_code(raw_text) or meta.get("school_code", "UNKNOWN")

    # Rule-based: cuba ekstrak skor numerik terus
    numeric_scores = _extract_numeric_scores(raw_text)

    # Kenal pasti bahagian kualitatif
    qualitative_parts = _extract_qualitative_parts(raw_text)

    conversion_method = "direct"
    ai_scores = {}

    if qualitative_parts:
        ai_scores = _ollama_convert(qualitative_parts, raw_text[:3000])
        if ai_scores:
            numeric_scores.update({k: v for k, v in ai_scores.items() if v is not None})
            conversion_method = "ai_converted" if not numeric_scores else "hybrid"

    # Kira confidence berdasarkan berapa banyak skor berjaya diekstrak
    confidence = _calc_confidence(numeric_scores, qualitative_parts, ai_scores)

    return {
        "school_code": school_code,
        "extracted_scores": numeric_scores,
        "qualitative_text": qualitative_parts[:500] if qualitative_parts else None,
        "data_type": _detect_data_type(numeric_scores, qualitative_parts),
        "conversion_method": conversion_method,
        "agent0_confidence": round(confidence, 3),
        "source_file": os.path.basename(file_path),
        "raw_text_preview": raw_text[:200],
    }


def _process_api_data(json_file: str, meta: dict) -> dict:
    """Process structured API data — sudah dalam bentuk numerik."""
    with open(json_file, "r") as f:
        data = json.load(f)

    school_code = data.get("school_code") or meta.get("school_code", "UNKNOWN")
    scores = data.get("scores", {})

    # Pastikan semua nilai adalah float
    numeric_scores = {}
    for k, v in scores.items():
        try:
            numeric_scores[k] = float(v)
        except (TypeError, ValueError):
            pass

    return {
        "school_code": school_code,
        "extracted_scores": numeric_scores,
        "qualitative_text": None,
        "data_type": "quantitative",
        "conversion_method": "direct",
        "agent0_confidence": 0.95 if numeric_scores else 0.10,
        "source_file": json_file,
        "raw_text_preview": json.dumps(scores)[:200],
    }


def _extract_school_code(text: str) -> Optional[str]:
    matches = SCHOOL_CODE_RE.findall(text)
    if matches:
        # Return first match, reconstruct code
        full_match = SCHOOL_CODE_RE.search(text)
        if full_match:
            return full_match.group(0).strip().upper()
    return None


def _extract_numeric_scores(text: str) -> dict:
    """Ekstrak skor numerik dari teks menggunakan regex."""
    scores = {}

    # Pattern: "Skor akademik: 78.5" atau "Academic Performance - 82"
    for match in HEADER_SCORE_RE.finditer(text):
        label = match.group(1).strip().lower()
        value = float(match.group(2))
        if 0 <= value <= 100:
            # Normalise label ke key
            key = re.sub(r'\s+', '_', label)[:40]
            scores[key] = value

    # Pattern: "85/100" atau "78.5%" dalam konteks
    for match in SCORE_RE.finditer(text):
        value = float(match.group(1))
        if 0 <= value <= 100:
            # Use position-based key jika tiada label
            if "unlabelled" not in scores:
                scores["unlabelled_score"] = value

    return scores


def _extract_qualitative_parts(text: str) -> str:
    """Cari ayat yang mengandungi penilaian kualitatif."""
    qual_keywords = list(QUAL_SCALE.keys()) + SKPMG2_DIMENSIONS
    lines = text.split('\n')
    qual_lines = []

    for line in lines:
        line_lower = line.lower()
        if any(kw in line_lower for kw in qual_keywords) and len(line.strip()) > 10:
            qual_lines.append(line.strip())

    return "\n".join(qual_lines[:30])  # Ambil max 30 baris


def _ollama_convert(qualitative_text: str, context: str) -> dict:
    """
    Hantar teks kualitatif ke Ollama llama3.2.
    Minta skor numerik 0-100 untuk dimensi SKPMG2.
    """
    prompt = f"""Anda adalah pakar penilaian sekolah Malaysia (SKPMG2).
Baca petikan laporan berikut dan jana skor numerik 0-100 untuk setiap dimensi yang disebut.

Dimensi yang boleh dinilai:
- academic_performance (prestasi akademik / pencapaian murid)
- discipline (disiplin / kelakuan murid)
- facilities (kemudahan / infrastruktur fizikal)
- leadership (kepimpinan pengetua / pengurusan)
- hygiene_canteen (kebersihan / kantin / sanitasi)
- co_curriculum (kokurikulum / aktiviti)
- safety (keselamatan persekitaran sekolah)

Teks laporan:
{qualitative_text}

Konteks tambahan:
{context[:500]}

Balas HANYA dalam format JSON (gunakan null jika dimensi tidak disebut):
{{"academic_performance": 75, "discipline": null, "facilities": 60,
  "leadership": null, "hygiene_canteen": null, "co_curriculum": null,
  "safety": null, "confidence": 0.80, "reasoning": "..."}}"""

    try:
        payload = json.dumps({
            "model": OLLAMA_MODEL,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "format": "json",
            "options": {"temperature": 0.0},
        }).encode()
        req = urllib.request.Request(
            f"{OLLAMA_URL}/api/chat",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        content = data["message"]["content"]
        content = re.sub(r'^```(?:json)?\s*', '', content.strip(), flags=re.IGNORECASE)
        content = re.sub(r'\s*```$', '', content.strip())
        result = json.loads(content)
        scores = {k: v for k, v in result.items()
                  if k not in ("confidence", "reasoning") and v is not None}
        return scores
    except Exception as e:
        logger.warning(f"Ollama conversion failed: {e}")
        return _rule_based_qual_convert(qualitative_text)


def _rule_based_qual_convert(text: str) -> dict:
    """Fallback: cari kata kunci kualitatif dan map ke skor."""
    text_lower = text.lower()
    scores = {}

    for phrase, score in QUAL_SCALE.items():
        if phrase in text_lower:
            # Find context around the phrase to determine dimension
            idx = text_lower.find(phrase)
            context = text_lower[max(0, idx-50):idx+50]

            for dim_kw, dim_key in [
                ("akademik", "academic_performance"), ("academic", "academic_performance"),
                ("disiplin", "discipline"), ("kemudahan", "facilities"),
                ("kantin", "hygiene_canteen"), ("kebersihan", "hygiene_canteen"),
            ]:
                if dim_kw in context:
                    scores[dim_key] = score
                    break

    return scores


def _detect_data_type(numeric_scores: dict, qualitative_text: str) -> str:
    has_numeric = bool(numeric_scores)
    has_qual = bool(qualitative_text and qualitative_text.strip())
    if has_numeric and has_qual:
        return "mixed"
    if has_qual:
        return "qualitative"
    return "quantitative"


def _calc_confidence(numeric_scores: dict, qual_text: str, ai_scores: dict) -> float:
    base = 0.0
    if numeric_scores:
        base += min(0.6, len(numeric_scores) * 0.15)
    if ai_scores:
        base += 0.25
    elif qual_text:
        base += 0.10
    return min(1.0, base)


def _error_result(msg: str, file_path: str) -> dict:
    return {
        "school_code": "UNKNOWN",
        "extracted_scores": {},
        "qualitative_text": None,
        "data_type": "unknown",
        "conversion_method": "error",
        "agent0_confidence": 0.0,
        "source_file": os.path.basename(file_path),
        "error": msg,
        "raw_text_preview": "",
    }


# ═══════════════════════════════════════════════════════════════════════════════
# BULK DOCUMENT — Multi-School Extraction (Pemeriksaan JN)
# ═══════════════════════════════════════════════════════════════════════════════

# ── Regex untuk kod sekolah yang lebih agresif ─────────────────────────────────
BULK_SCHOOL_CODE_RE = re.compile(
    r'\b((?:SK|SMK|SJK|MRSM|SBP|SKK|SEKOLAH)\s*(?:[A-Za-z]+\s*)*\d{3,})\b',
    re.IGNORECASE,
)

# ── Section headers yang menandakan permulaan blok sekolah baharu ──────────────
SCHOOL_SECTION_HEADERS = [
    r'^.*(?:SEKOLAH|SCHOOL)\s*[:\-]?\s*',
    r'^.*KOD\s*SEKOLAH\s*[:\-]?\s*',
    r'^.*NAMA\s*SEKOLAH\s*[:\-]?\s*',
    r'^\d+[\.\)]\s*(?:SEKOLAH|SCHOOL)?\s*',
    r'^={2,}\s*$',  # separator line
]


def find_all_school_codes(text: str) -> list[str]:
    """
    Cari SEMUA kod sekolah dalam dokumen bulk Pemeriksaan JN.
    Returns deduplicated list mengikut urutan penemuan.
    """
    seen = set()
    codes = []
    for match in BULK_SCHOOL_CODE_RE.finditer(text):
        code = match.group(1).strip().upper()
        # Normalise: buang extra whitespace, standardise prefix
        code = re.sub(r'\s+', '', code)
        if code not in seen and len(code) >= 4:
            seen.add(code)
            codes.append(code)
    return codes


def split_document_by_school(raw_text: str, school_codes: list[str]) -> dict[str, str]:
    """
    Pecahkan dokumen bulk kepada bahagian per sekolah.
    Cuba beberapa strategi:
      1. Cari section header dengan nama/kod sekolah
      2. Fallback: bahagi ikut paragraph besar yang mengandungi kod
    Returns {school_code: section_text}
    """
    if not school_codes:
        return {}

    sections = {}
    lines = raw_text.split('\n')

    # Strategi 1: Cari baris yang mengandungi kod sekolah sebagai section marker
    code_positions = []
    for i, line in enumerate(lines):
        line_upper = line.upper().replace(' ', '')
        for code in school_codes:
            code_compact = code.replace(' ', '')
            if code_compact in line_upper:
                code_positions.append((i, code))
                break

    if len(code_positions) >= 2:
        # Ada section markers — split antara markers
        for idx, (pos, code) in enumerate(code_positions):
            start = pos
            end = code_positions[idx + 1][0] if idx + 1 < len(code_positions) else len(lines)
            section = '\n'.join(lines[start:end]).strip()
            if section:
                sections[code] = section
        return sections

    # Strategi 2: Fallback — beri keseluruhan dokumen kepada setiap sekolah
    # (Agent 0 akan ekstrak skor spesifik dari keseluruhan teks)
    for code in school_codes:
        sections[code] = raw_text

    return sections


def _ai_detect_schools_and_scores(raw_text: str) -> list[dict]:
    """
    [NEW] Gunakan LLM (Ollama/DeepSeek) untuk membaca dokumen Pemeriksaan JN
    dan mengenal pasti:
      - Senarai sekolah yang disebut
      - Skor/syor untuk setiap standard SKPM per sekolah

    Digunakan apabila regex gagal atau untuk dokumen kompleks.
    Returns: [{"school_code": "SBP001", "domain_scores": {...}, "syor_text": "..."}, ...]
    """
    prompt = f"""Anda adalah sistem OCR pintar untuk dokumen Pemeriksaan Jemaah Nazir (JN) Malaysia.
Baca petikan dokumen berikut dan ekstrak:

1. Senarai KOD SEKOLAH (format: SK..., SMK..., SBP..., MRSM..., SJK..., SKK... diikuti nombor)
2. Untuk SETIAP sekolah, cari:
   - Skor atau syor untuk standard SKPM: Kekuatan, Kepimpinan, Pengurusan Organisasi,
     Pengurusan Kurikulum, Pengurusan Kokurikulum, Pengurusan HEM, PdPc, Kemenjadian Murid
   - Jika ada skor numerik (0-100), ekstrak nilainya
   - Jika ada pernyataan syor (kualitatif), petik teks asal

Dokumen (mungkin panjang — fokus pada bahagian yang mengandungi kod sekolah dan skor):
{raw_text[:8000]}

Balas HANYA dalam format JSON:
{{"schools": [
  {{"school_code": "SBP001", "scores": {{"kepimpinan": 85.0, "pdpc": 78.5}}, "syor": ["Syor 1...", "Syor 2..."]}},
  ...
], "inspection_theme": "PAJSK/Kurikulum/PPD", "confidence": 0.85}}"""

    try:
        payload = json.dumps({
            "model": OLLAMA_MODEL,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "format": "json",
            "options": {"temperature": 0.0, "num_ctx": 8192},
        }).encode()
        req = urllib.request.Request(
            f"{OLLAMA_URL}/api/chat",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read())
        content = data["message"]["content"]
        content = re.sub(r'^```(?:json)?\s*', '', content.strip(), flags=re.IGNORECASE)
        content = re.sub(r'\s*```$', '', content.strip())
        result = json.loads(content)
        return result.get("schools", [])
    except Exception as e:
        logger.warning(f"AI school detection failed: {e}")
        return []


def run_bulk(file_path: str, source_type: str,
             school_codes_filter: Optional[list[str]] = None,
             source_metadata: Optional[dict] = None) -> list[dict]:
    """
    [NEW] Proses satu dokumen bulk Pemeriksaan JN yang mengandungi
    PELBAGAI sekolah. Menyahserialkan kepada senarai keputusan per sekolah.

    Flow:
      1. Parse dokumen → raw_text
      2. Cari SEMUA kod sekolah (find_all_school_codes)
      3. Cuba split dokumen ikut section sekolah
      4. Untuk setiap sekolah, proses extracted_scores & qualitative → skor
      5. Opsyenal: tapis dengan school_codes_filter

    Returns: list of dict (sama format seperti run() tapi untuk multiple schools)
    """
    source_metadata = source_metadata or {}
    raw_text = _parse_document(file_path, source_type)

    # Langkah 1: Cari semua kod sekolah dalam dokumen
    detected_codes = find_all_school_codes(raw_text)

    # Langkah 2: Jika regex tak jumpa, guna AI detection
    if not detected_codes and len(raw_text) > 200:
        logger.info("Regex failed to find school codes — trying AI detection...")
        ai_schools = _ai_detect_schools_and_scores(raw_text)
        if ai_schools:
            return _build_bulk_results_from_ai(ai_schools, raw_text, file_path, source_metadata)

    # Langkah 3: Jika masih tak jumpa, fallback — cuba dari metadata
    if not detected_codes:
        fallback_code = source_metadata.get("school_code", "UNKNOWN")
        detected_codes = [fallback_code] if fallback_code != "UNKNOWN" else []

    if not detected_codes:
        logger.warning(f"No school codes found in document: {file_path}")
        return []

    # Langkah 4: Tapis ikut filter jika ada
    if school_codes_filter:
        detected_codes = [c for c in detected_codes if c in school_codes_filter]

    # Langkah 5: Split dokumen kepada bahagian per sekolah
    sections = split_document_by_school(raw_text, detected_codes)

    # Langkah 6: Proses setiap bahagian
    results = []
    for code in detected_codes:
        section_text = sections.get(code, raw_text)
        # Guna _process_text dengan section spesifik
        result = _process_text(section_text, source_type, file_path,
                               {**source_metadata, "school_code": code, "is_bulk": True})
        result["school_code"] = code  # override
        results.append(result)

    return results


def _build_bulk_results_from_ai(ai_schools: list[dict], raw_text: str,
                                 file_path: str, meta: dict) -> list[dict]:
    """Bina keputusan bulk dari output AI detection."""
    results = []
    for s in ai_schools:
        code = s.get("school_code", "UNKNOWN")
        scores = s.get("scores", {})
        syor_lines = s.get("syor", [])

        # Normalise scores
        numeric_scores = {}
        for k, v in scores.items():
            try:
                numeric_scores[k] = float(v)
            except (TypeError, ValueError):
                pass

        results.append({
            "school_code": code,
            "extracted_scores": numeric_scores,
            "qualitative_text": "\n".join(syor_lines) if syor_lines else raw_text[:500],
            "data_type": "mixed" if syor_lines else "quantitative",
            "conversion_method": "ai_converted",
            "agent0_confidence": s.get("confidence", 0.80),
            "source_file": os.path.basename(file_path),
            "raw_text_preview": raw_text[:200],
        })

    return results
