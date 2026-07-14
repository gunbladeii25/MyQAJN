"""
Create a dummy Pemeriksaan JN DOCX file with MULTIPLE school codes
embedded in content — NOT in filename — and upload to Google Drive.

Usage: python scripts/create_dummy_inspection.py
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2 import service_account
from dotenv import load_dotenv
load_dotenv()

FOLDER_ID = "1manDr6M9vF6eNLhQZvYIVdg15FcdA7EO"
SA_KEY = os.getenv("GDRIVE_SERVICE_ACCOUNT_JSON", os.path.join(os.path.dirname(__file__), "..", "secrets", "gdrive-sa.json"))

# ── Dummy DOCX content ──────────────────────────────────────────────────────
DOC_CONTENT = """
PEMERIKSAAN PAJSK 2026
KEMENTERIAN PENDIDIKAN MALAYSIA
JEMAAH NAZIR
Tarikh Pemeriksaan: 12 Julai 2026
Tema: Pendidikan Jasmani & Sukan (PAJSK)

LAPORAN SYOR PEMERIKSAAN

--------------------------------------------------------------------------------
SEKOLAH: MRSM Kuala Terengganu
KOD SEKOLAH: MRSM001
JENIS: MRSM | NEGERI: Terengganu
--------------------------------------------------------------------------------

Standard 1 — Kepimpinan:
  Skor: 82.0/100
  Syor: PGB perlu memperkukuh peranan sebagai peneraju instruksional dengan 
  pemantauan PdP secara berkala. Program mentoring guru baharu wajar diperluaskan.

Standard 2 — Pengurusan Organisasi:
  Skor: 85.5/100
  Syor: Pengurusan sumber manusia dan kewangan sekolah hendaklah diperkemas 
  selaras dengan tatacara semasa KPM.

Standard 3.2 — Pengurusan Kokurikulum (PAJSK):
  Skor: 78.0/100
  Syor: Penyertaan murid dalam aktiviti sukan dan kokurikulum perlu diperluas 
  ke peringkat daerah dan negeri. Kemudahan padang permainan wajar dinaik taraf.

Standard 4 — PdPc:
  Skor: 76.5/100
  Syor: Kualiti pembelajaran dan pemudahcaraan guru PJ perlu ditingkatkan 
  melalui bimbingan instruksional berfokus.

Standard 5 — Kemenjadian Murid:
  Skor: 88.0/100
  Syor: Program peningkatan sahsiah dan disiplin murid melalui aktiviti 
  kokurikulum menunjukkan kesan positif.

--------------------------------------------------------------------------------
SEKOLAH: Sekolah Berasrama Penuh Integrasi Gombak
KOD SEKOLAH: SBP001
JENIS: SBP | NEGERI: Selangor
--------------------------------------------------------------------------------

Standard A — Kekuatan (Prasarana & Sumber):
  Skor: 93.5/100
  Syor: Prasarana dan kemudahan fizikal pada tahap cemerlang. Penyelenggaraan 
  berkala wajar diteruskan untuk memastikan kelestarian fasiliti.

Standard 1 — Kepimpinan:
  Skor: 90.0/100
  Syor: Kepimpinan PGB sangat baik. Program succession planning untuk 
  barisan pemimpin pertengahan wajar dimulakan.

Standard 3.1 — Pengurusan Kurikulum:
  Skor: 87.5/100
  Syor: Pelaksanaan kurikulum berkesan dengan pencapaian akademik memberangsangkan.
  Program pengayaan untuk murid cemerlang wajar diperkenalkan.

Standard 5 — Kemenjadian Murid:
  Skor: 92.0/100
  Syor: Pencapaian akademik dan sahsiah murid cemerlang. Program outreach 
  ke sekolah sekitar boleh dijadikan model rujukan.

--------------------------------------------------------------------------------
SEKOLAH: Sekolah Menengah Kebangsaan Taman Jaya
KOD SEKOLAH: SMK001
JENIS: SMK | NEGERI: Selangor
--------------------------------------------------------------------------------

Standard A — Kekuatan (Prasarana & Sumber):
  Skor: 62.0/100
  Syor: Kemudahan asas seperti bilik darjah dan tandas memerlukan naik taraf 
  segera. Peruntukan penyelenggaraan wajar disalurkan.

Standard 1 — Kepimpinan:
  Skor: 70.5/100
  Syor: Pengurusan sekolah pada tahap sederhana. PGB wajar diberi latihan 
  kepimpinan dan pengurusan perubahan.

Standard 3.3 — Pengurusan Hal Ehwal Murid (HEM):
  Skor: 65.0/100
  Syor: Disiplin murid perlu pemantauan lebih kerap. Sistem intervensi awal 
  untuk murid berisiko wajar diperkukuh.

Standard 4 — PdPc:
  Skor: 68.0/100
  Syor: Kualiti pengajaran memerlukan bimbingan berterusan. Program 
  professional learning community (PLC) wajar diaktifkan semula.

================================================================================
DISEDIAKAN OLEH:
Pasukan Pemeriksaan Jemaah Nazir
Tarikh: 12 Julai 2026
================================================================================
"""


def create_docx(content: str, output_path: str):
    """Create a .docx file from plain text content."""
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = 11 * 12700  # 11pt in EMU

    for line in content.strip().split('\n'):
        line = line.rstrip()
        if not line:
            doc.add_paragraph('')
        elif line.startswith('===') or line.startswith('---'):
            p = doc.add_paragraph()
            p.alignment = 1  # center
            run = p.add_run(line.strip('-=').strip())
            run.bold = True
        elif line.startswith('SEKOLAH:') or line.startswith('KOD SEKOLAH:'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif line.startswith('Standard') or line.startswith('Skor:'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif line.startswith('  Syor:'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = 360000  # ~1cm
            p.add_run(line.strip())
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"✅ DOCX created: {output_path}")


def upload_to_gdrive(file_path: str, folder_id: str, sa_key_path: str):
    """Upload file to Google Drive folder."""
    SCOPES = ['https://www.googleapis.com/auth/drive.file']
    creds = service_account.Credentials.from_service_account_file(sa_key_path, scopes=SCOPES)
    service = build('drive', 'v3', credentials=creds)

    file_name = os.path.basename(file_path)
    mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    file_metadata = {
        'name': file_name,
        'parents': [folder_id],
    }
    media = MediaFileUpload(file_path, mimetype=mime_type, resumable=True)

    uploaded = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id,name,size,webViewLink',
    ).execute()

    print(f"✅ Uploaded to GDrive: {uploaded.get('name')}")
    print(f"   File ID: {uploaded.get('id')}")
    print(f"   Link: {uploaded.get('webViewLink')}")
    return uploaded


if __name__ == '__main__':
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'data')
    os.makedirs(output_dir, exist_ok=True)
    docx_path = os.path.join(output_dir, 'Pemeriksaan_PAJSK_2026_Bulk.docx')

    create_docx(DOC_CONTENT, docx_path)

    # Also create a second themed file for variety
    docx_path2 = os.path.join(output_dir, 'Pemeriksaan_Kurikulum_2026_Bulk.docx')
    kurikulum_content = DOC_CONTENT.replace('PAJSK', 'Kurikulum').replace(
        'Pendidikan Jasmani & Sukan (PAJSK)', 'Kurikulum & Akademik'
    )
    create_docx(kurikulum_content, docx_path2)

    # Upload both to GDrive
    print("\n📤 Uploading to Google Drive...")
    try:
        upload_to_gdrive(docx_path, FOLDER_ID, SA_KEY)
        upload_to_gdrive(docx_path2, FOLDER_ID, SA_KEY)
        print("\n🎉 Done! Both files uploaded.")
        print("   Now the AI can detect schools from file CONTENT, not filename.")
    except Exception as e:
        print(f"\n⚠️  Upload failed: {e}")
        print("   Files created locally — upload manually if needed.")
        print(f"   {docx_path}")
        print(f"   {docx_path2}")
